"""
Módulo para generar el informe en formato DOCX
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
from textos import (
    PARRAFOS_FIJOS, PARRAFOS_PT, TEXTO_ERROR_y_INT, TEXTO_FINAL_PosibleDislexia
)

def agregar_portada(doc: Document, nombre_completo: str, datos: dict) -> None:
    """
    Agrega la portada del informe
    
    Args:
        doc: Documento de Word
        nombre_completo: Nombre completo del encuestado
        datos: Diccionario con datos generales (edad, fecha_aplicacion, etc.)
    """
    # Título
    titulo = doc.add_heading('PRUEBA STROOP, TAREA DE INTERFERENCIA PALABRA-COLOR', 0)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Aumentar tamaño de fuente del título
    for run in titulo.runs:
        run.font.size = Pt(24)
    
    # Espacio
    doc.add_paragraph().add_run().add_break()
    
    # Información del participante
    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Nombre del encuestado
    nombre_run = info.add_run(f"Nombre del encuestado: {nombre_completo}\n")
    nombre_run.bold = True
    nombre_run.font.size = Pt(14)
    
    # Edad
    if datos.get('edad'):
        edad_run = info.add_run(f"Edad: {datos['edad']} años\n")
        edad_run.font.size = Pt(14)
    
    # Fecha de aplicación (si está disponible)
    if datos.get('fecha_aplicacion'):
        fecha_app_run = info.add_run(f"Fecha de aplicación: {datos['fecha_aplicacion']}\n")
        fecha_app_run.font.size = Pt(14)
    
    # Fecha del informe
    fecha_actual = datetime.now().strftime("%d/%m/%Y")
    fecha_informe_run = info.add_run(f"Fecha del informe: {fecha_actual}\n")
    fecha_informe_run.font.size = Pt(14)
    
    # Espacio
    doc.add_paragraph().add_run().add_break()
    
    # Nota confidencial
    nota = doc.add_paragraph()
    nota.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nota_run = nota.add_run(f"Este es un informe de evaluación cognitiva, obtenido a partir del rendimiento de {nombre_completo} en la prueba Stroop (Tarea de Interferencia Palabra-Color).")
    doc.add_paragraph()  # Espacio
    nota.add_run("\n" + "-" * 50 + "\n")  # Línea de separación
    nota.add_run("\nInforme confidencial  de carácter educativo u orientativo. No es un diagnóstico clínico y su interpretación es conveniente la realice un profesional competente.")
    nota_run.italic = True
    nota_run.font.size = Pt(12)

def crear_informe_docx(resultados, clasificaciones, nombre_caso="caso", scale_factor_width=0.8, scale_factor_height=0.5):
    """
    Crea el documento Word con el informe del test Stroop
    
    Args:
        resultados: Dict con puntuaciones directas y datos del encuestado
        clasificaciones: Dict con clasificaciones (PT)
        nombre_caso: Nombre del evaluado (fallback si no está en resultados)
        
    Returns:
        Document: Documento Word generado
    """
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Extraer nombres de resultados
    nombre_completo = resultados.get('nombre_completo') or nombre_caso
    nombre = resultados.get('nombre') or nombre_caso
    
    # Preparar datos para la portada
    datos_portada = {
        'edad': resultados.get('edad'),
        'fecha_aplicacion': resultados.get('fecha_aplicacion')
    }
    
    # 1. PORTADA (Primera página)
    agregar_portada(doc, nombre_completo, datos_portada)
    doc.add_page_break()
    
    # ========================================================================
    # TÍTULO E INTRODUCCIÓN
    # ========================================================================
    # ========================================================================
    # DESCRIPCIÓN DE LA PRUEBA
    # ========================================================================
    
    parrafo = doc.add_paragraph()
    run = parrafo.add_run(PARRAFOS_FIJOS['titulo_general_prueba'])
    run.bold = True
    doc.add_paragraph(PARRAFOS_FIJOS['objetivo_prueba'])

    parrafo = doc.add_paragraph()
    run = parrafo.add_run(PARRAFOS_FIJOS['titulo_procedimiento'])
    run.bold = True
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_procedimiento1'])
    
    # Imagen parte P
    script_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'images')
    grafico_path = os.path.join(script_dir, 'P.png')
    if os.path.exists(grafico_path):
        try:
            # Obtener dimensiones de la imagen
            img = Image.open(grafico_path)
            width, height = img.size
            # Aplicar factor de escala
            new_width = Inches(width / 96 * scale_factor_width)  # 96 DPI
            new_height = Inches(height / 96 * scale_factor_height)
            # Agregar imagen
            doc.add_picture(grafico_path, width=new_width, height=new_height)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            print(f"Advertencia: No se pudo insertar la imagen P.png: {e}")
    
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_procedimiento2'])
    
    # Imagen parte C
    grafico_path = os.path.join(script_dir, 'C.png')
    if os.path.exists(grafico_path):
        try:
            img = Image.open(grafico_path)
            width, height = img.size
            new_width = Inches(width / 96 * scale_factor_width)
            new_height = Inches(height / 96 * scale_factor_height)
            doc.add_picture(grafico_path, width=new_width, height=new_height)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            print(f"Advertencia: No se pudo insertar la imagen C.png: {e}")
    
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_procedimiento3'])
    
    # Imagen parte PC
    grafico_path = os.path.join(script_dir, 'PC.png')
    if os.path.exists(grafico_path):
        try:
            img = Image.open(grafico_path)
            width, height = img.size
            new_width = Inches(width / 96 * scale_factor_width)
            new_height = Inches(height / 96 * scale_factor_height)
            doc.add_picture(grafico_path, width=new_width, height=new_height)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            print(f"Advertencia: No se pudo insertar la imagen PC.png: {e}")

    parrafo = doc.add_paragraph()
    run = parrafo.add_run(PARRAFOS_FIJOS['titulo_indices'])
    run.bold = True
    doc.add_paragraph(PARRAFOS_FIJOS['descripcion_indices'])
    doc.add_page_break()

    # ========================================================================
    # INSERTAR Tabla resultados
    # ========================================================================
    
    # Título de resultados
    titulo_resultados = doc.add_paragraph()
    titulo_resultados.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo_resultados.add_run(PARRAFOS_FIJOS['titulo_resultados'].format(nombre_completo=nombre_completo))
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()  # Espacio
    
    doc.add_paragraph(PARRAFOS_FIJOS['texto_tabla_resultados'].format(nombre=nombre, nombre_completo=nombre_completo))

    # Tabla PDs, PTs y clasificaciones
    tabla = doc.add_table(rows=4, cols=6)

    # Aplicar estilo simple con bordes negros y encabezado con fondo gris claro
    tabla.style = 'Table Grid'
    # Aplicar fondo gris claro al encabezado
    for cell in tabla.rows[0].cells:
        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))
    
    # Encabezados
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Palabras (P)'
    hdr_cells[2].text = 'Colores (C)'
    hdr_cells[3].text = 'Palabras-Colores (PC)'
    hdr_cells[4].text = 'Interferencia (INT)'
    hdr_cells[5].text = 'Errores (E)'

    # Centrar el texto en las celdas del encabezado
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
    # Fila 1: PD (Puntuación Directa)
    row1 = tabla.rows[1].cells
    row1[0].text = 'Puntuación Directa (PD)'
    row1[1].text = str(resultados.get('PD_P', 0))
    row1[2].text = str(resultados.get('PD_C', 0))
    row1[3].text = str(resultados.get('PD_PC', 0))
    row1[4].text = str(round(resultados.get('Indice_interferencia', 0), 2))
    row1[5].text = str(resultados.get('PD_E', 0))
    
    # Fila 2: PT (Puntuación Típica) - E no tiene PT
    row2 = tabla.rows[2].cells
    row2[0].text = 'Puntuación Típica (PT)'
    row2[1].text = str(resultados.get('PT_P', '-'))
    row2[2].text = str(resultados.get('PT_C', '-'))
    row2[3].text = str(resultados.get('PT_PC', '-'))
    row2[4].text = str(resultados.get('PT_INT', '-'))
    row2[5].text = '-'  # E no tiene PT
    
    # Fila 3: Clasificación (bajo/normal/alto)
    row3 = tabla.rows[3].cells
    row3[0].text = 'Rendimiento'
    row3[1].text = str(resultados.get('Clasificacion_P', '-'))
    row3[2].text = str(resultados.get('Clasificacion_C', '-'))
    row3[3].text = str(resultados.get('Clasificacion_PC', '-'))
    row3[4].text = str(resultados.get('Clasificacion_INT', '-'))
    row3[5].text = str(resultados.get('Clasificacion_E', '-'))

    # Verde si es alto, rojo si es bajo (columnas: Palabras, Colores, Palabras-Colores e Interferencia)
    green_if_high = [1, 2, 3, 4]  # Índices de columnas para P, C, PC e INT
    # Rojo si es alto, verde si es bajo (columnas: Errores)
    red_if_high = [5]
    
    for idx in green_if_high:
        clasificacion = str(row3[idx].text).lower()
        if clasificacion == 'alto':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="90EE90"/>'.format(nsdecls('w'))))
        elif clasificacion == 'bajo':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF6B6B"/>'.format(nsdecls('w'))))
    
    for idx in red_if_high:
        clasificacion = str(row3[idx].text).lower()
        if clasificacion == 'alto':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF6B6B"/>'.format(nsdecls('w'))))
        elif clasificacion == 'bajo':
            row3[idx]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="90EE90"/>'.format(nsdecls('w'))))


    # Centrar el texto en las celdas de las filas 1, 2 y 3
    for row in [row1, row2, row3]:
        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Espacio

    # ========================================================================
    # PÁRRAFOS CONDICIONALES
    # ========================================================================

    # Construir clave para buscar el párrafo condicional de P, C y PC
    clas_P = resultados.get('Clasificacion_P', 'normal')
    clas_C = resultados.get('Clasificacion_C', 'normal')
    clas_PC = resultados.get('Clasificacion_PC', 'normal')
    
    # Función auxiliar para buscar la clave correcta en el diccionario
    def encontrar_clave_PT(p, c, pc):
        """Genera posibles claves y busca en PARRAFOS_PT"""
        # Probar diferentes variantes de clave
        claves_posibles = []
        
        # Caso especial: todos iguales
        if p == c == pc:
            if p == 'normal':
                claves_posibles.append('P, C y PC normales')
            elif p == 'bajo':
                claves_posibles.extend(['P, C y PC bajo', 'P, C y PC bajo '])
            else:  # alto
                claves_posibles.extend(['P, C y PC alto', 'P, C y PC alto '])
        # Dos iguales: P y C (pero PC diferente)
        elif p == c and p != pc:
            claves_posibles.extend([
                f'P y C {p} y PC {pc}',
                f'P y C {p} y PC {pc} ',
                f'P C {p} y PC {pc}',  # Variante sin 'y' al principio
                f'P C {p} y PC {pc} ',
            ])
        # Dos iguales: C y PC (pero P diferente)
        elif c == pc and c != p:
            claves_posibles.extend([
                f'P {p} y C y PC {c}',
                f'P {p} y C y PC {c} ',
                f'P {p}, C y PC {c}',
                f'P {p}, C y PC {c} ',
            ])
        # Todos diferentes O P y PC iguales (C diferente)
        # En textos.py, P=PC se escribe como si todos fueran diferentes
        else:
            claves_posibles.extend([
                f'P {p}, C {c} y PC {pc}',
                f'P {p}, C {c} y PC {pc} ',
                f'P {p} y C {c} y PC {pc}',
                f'P {p} y C {c} y PC {pc} ',
            ])
        
        # Buscar primera clave que exista
        for clave in claves_posibles:
            if clave in PARRAFOS_PT:
                return clave
        
        return None
    
    # Buscar y añadir párrafo correspondiente
    texto_condicional = ""
    clave_PT = encontrar_clave_PT(clas_P, clas_C, clas_PC)
    
    if clave_PT and clave_PT in PARRAFOS_PT:
        texto_pt = PARRAFOS_PT[clave_PT].format(nombre=nombre)
        texto_condicional += texto_pt
    else:
        # Si no se encuentra la clave, agregar mensaje de debug
        print(f"Advertencia: No se encontró texto para P={clas_P}, C={clas_C}, PC={clas_PC}")
        texto_condicional += f"Resultados: P={clas_P}, C={clas_C}, PC={clas_PC}. "

    # Párrafo condicional de E (Errores) e INT (Interferencia)
    clas_E = resultados.get('Clasificacion_E', 'normal')
    clas_INT = resultados.get('Clasificacion_INT', 'normal')
    
    # Probar diferentes variantes de clave
    claves_e_int_posibles = [
        f'E {clas_E} y INT {clas_INT}',
        f'E y INT {clas_INT}',  # Caso especial para 'E y INT bajo'
    ]
    
    texto_encontrado = False
    for clave_E_INT in claves_e_int_posibles:
        if clave_E_INT in TEXTO_ERROR_y_INT:
            texto_e_int = TEXTO_ERROR_y_INT[clave_E_INT].format(nombre=nombre)
            texto_condicional += texto_e_int
            texto_encontrado = True
            break
    
    if not texto_encontrado:
        print(f"Advertencia: No se encontró texto para E={clas_E}, INT={clas_INT}")
    
    # Párrafo final sobre posible dislexia (solo en ciertos casos)
    if clave_PT and clave_PT in TEXTO_FINAL_PosibleDislexia:
        texto_final = TEXTO_FINAL_PosibleDislexia[clave_PT].format(nombre=nombre)
        texto_condicional += texto_final

    # Añadir todo el texto condicional al documento en un solo párrafo
    doc.add_paragraph(texto_condicional)

    # ========================================================================
    return doc


def guardar_informe(doc, ruta_salida):
    """
    Guarda el documento generado
    
    Args:
        doc: Documento Word
        ruta_salida: Ruta donde guardar el archivo
    """
    doc.save(ruta_salida)
    print(f"Informe generado exitosamente en: {ruta_salida}")
